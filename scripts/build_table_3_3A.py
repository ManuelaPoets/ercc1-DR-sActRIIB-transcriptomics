"""Build Table 3.3A (top driver genes) as a REAL Word table + caption,
replicating notebook cell 43's logic. Paste-ready docx for the Drive master."""
import pandas as pd, glob, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

BASE = r'D:\manuela\Thesis\results'
CATS = [  # label, file-suffix, restoration column, driving-intervention LFC column
    ('DR-exclusive',     'intervention_impact_comparison_DR.csv',                'Restoration_DR',      'LFC_DR'),
    ('sActRIIB-exclusive','intervention_impact_comparison_sAct.csv',             'Restoration_sAct',    'LFC_sAct'),
    ('Combination-only', 'intervention_impact_comparison_only_rescued_in_combi.csv','Restoration_Combined','LFC_Combined'),
]
TISSUES = [('Kidney', 'Kidney_analysis_results', 'kidney', False),
           ('Muscle', 'Muscle_v2_analysis_results', 'muscle_v2', True)]  # drop_sact=True for muscle
HEADERS = ['Category', 'Gene', 'Aging log₂FC', 'DR log₂FC', 'sAct log₂FC', 'Combi log₂FC', 'Restoration']

def fmt(v):
    v = float(v)
    return '0.00' if abs(round(v, 2)) < 0.005 else f"{v:+.2f}"   # no signed-zero

def rows_for(folder, tissue, drop_sact=False):
    out = []
    for label, suf, restcol, lfccol in CATS:
        if drop_sact and label.startswith('sActRIIB'):
            continue                                  # no genuine sActRIIB-exclusive genes in muscle
        f = glob.glob(os.path.join(BASE, folder, '**', f'{tissue}_{suf}'), recursive=True)
        df = pd.read_csv(f[0]).dropna(subset=['Symbol', restcol])  # named genes w/ a restoration score
        import numpy as np
        df = df[np.isfinite(df[restcol])]
        # rank by the magnitude of THIS category's driving-intervention LFC
        df = df.reindex(df[lfccol].abs().sort_values(ascending=False).index).head(5)
        for _, g in df.iterrows():
            out.append((label, str(g['Symbol']),
                        fmt(g['LFC_Age']), fmt(g['LFC_DR']),
                        fmt(g['LFC_sAct']), fmt(g['LFC_Combined']),
                        f"{float(g[restcol]):.2f}"))
    return out

doc = Document()
# caption
cap = doc.add_paragraph()
r = cap.add_run('Table 3.3A. ')
r.bold = True
cap.add_run('Top rescued-to-normal genes per intervention-driver category in kidney and muscle. '
            'The five genes most strongly shifted by the driving intervention (largest |intervention '
            'log₂FC|) are listed for each category. Aging, DR, '
            'sActRIIB and combined log₂ fold-changes are from DESeq2 (n = 3 per condition). '
            'Restoration = intervention log₂FC / |aging log₂FC| (1 = full reversal). In muscle, no genes were rescued exclusively by sActRIIB above noise level, so that category is omitted (see Section 3.3).')

from itertools import groupby
for title, folder, tissue, drop_sact in TISSUES:
    h = doc.add_paragraph(); hr = h.add_run(title); hr.bold = True; hr.font.size = Pt(11)
    data = rows_for(folder, tissue, drop_sact)
    t = doc.add_table(rows=1 + len(data), cols=len(HEADERS))
    t.style = 'Table Grid'
    for j, htxt in enumerate(HEADERS):
        c = t.cell(0, j); c.text = htxt
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            if j == 0: continue                      # category set once after merge
            c = t.cell(i, j); c.text = val
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if j == 1: c.paragraphs[0].runs[0].italic = True   # gene symbols italic
    # vertical-merge the Category column per consecutive group (any group size/count)
    ridx = 1
    for cat, grp in groupby(data, key=lambda r: r[0]):
        n = len(list(grp))
        m = t.cell(ridx, 0)
        for k in range(ridx + 1, ridx + n):
            m = m.merge(t.cell(k, 0))
        m.text = cat
        m.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        m.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        m.paragraphs[0].runs[0].font.size = Pt(9)
        m.paragraphs[0].runs[0].bold = True
        ridx += n

out = r'D:\manuela\Thesis\analysis\Table_3_3A_driver_genes.docx'
doc.save(out)
print('saved', out)
# echo the data for review
for title, folder, tissue, drop_sact in TISSUES:
    print(f'\n=== {title} ===')
    for row in rows_for(folder, tissue, drop_sact):
        print('  ', row)
